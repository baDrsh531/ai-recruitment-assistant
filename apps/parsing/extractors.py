"""Extraction du contenu brut d'un CV (PDF ou DOCX).

Cette couche ne fait aucune interpretation : elle produit du texte, des
coordonnees de mots et, au besoin, des images de pages. C'est le socle sur
lequel l'ancrage des preuves (`evidence.py`) devient possible : sans les
positions des mots, impossible de surligner un passage dans le PDF.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field

import fitz  # PyMuPDF
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# Rendu des pages pour le modele vision.
#
# 150 dpi paraissait suffisant — c'est faux, et le harnais d'evaluation l'a
# montre : sur un CV scanne, Qwen3-VL n'y lisait que les titres de sections
# (« COMPETENCES », « LANGUES ») au lieu des competences elles-memes. A 220 dpi
# il les lit toutes. Le cout est reel — 85 Ko contre 272 Ko par page, 11,6 s
# contre 15,4 s — mais une extraction fausse ne vaut rien, quel qu'en soit le prix.
RENDER_DPI = 220
MAX_RENDERED_PAGES = 4


class ExtractionError(RuntimeError):
    pass


@dataclass(slots=True)
class Word:
    """Un mot avec sa position dans la page et son offset dans le texte reconstruit."""

    text: str
    bbox: tuple[float, float, float, float]
    char_start: int
    char_end: int


@dataclass(slots=True)
class Page:
    number: int  # 1-indexe
    text: str
    words: list[Word] = field(default_factory=list)
    width: float = 0.0
    height: float = 0.0

    @property
    def char_count(self) -> int:
        return len(self.text.strip())


@dataclass(slots=True)
class ExtractedDocument:
    pages: list[Page]
    source: str  # "pdf" | "docx"

    @property
    def page_count(self) -> int:
        return len(self.pages)

    @property
    def full_text(self) -> str:
        return "\n\n".join(
            f"--- page {page.number} ---\n{page.text}" for page in self.pages
        )

    @property
    def char_count(self) -> int:
        return sum(page.char_count for page in self.pages)


# --- PDF -------------------------------------------------------------------
def extract_pdf(data: bytes) -> ExtractedDocument:
    """Extrait le texte natif d'un PDF avec les positions de chaque mot.

    Le texte de la page est reconstruit a partir de la liste de mots plutot
    que via `get_text("text")` : c'est la seule facon de garantir qu'un offset
    dans le texte correspond exactement a un mot dont on connait la bbox.
    """
    try:
        document = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"PDF illisible : {exc}") from exc

    pages: list[Page] = []
    with document:
        for index, page in enumerate(document, start=1):
            # (x0, y0, x1, y1, mot, bloc, ligne, num_mot)
            raw_words = page.get_text("words")
            raw_words.sort(key=lambda w: (w[5], w[6], w[7]))

            buffer: list[str] = []
            words: list[Word] = []
            cursor = 0
            previous_line: tuple[int, int] | None = None

            for x0, y0, x1, y1, token, block_no, line_no, _ in raw_words:
                line_key = (block_no, line_no)
                if previous_line is None:
                    separator = ""
                elif line_key != previous_line:
                    separator = "\n\n" if block_no != previous_line[0] else "\n"
                else:
                    separator = " "

                buffer.append(separator)
                cursor += len(separator)

                buffer.append(token)
                words.append(
                    Word(
                        text=token,
                        bbox=(float(x0), float(y0), float(x1), float(y1)),
                        char_start=cursor,
                        char_end=cursor + len(token),
                    )
                )
                cursor += len(token)
                previous_line = line_key

            pages.append(
                Page(
                    number=index,
                    text="".join(buffer),
                    words=words,
                    width=float(page.rect.width),
                    height=float(page.rect.height),
                )
            )

    return ExtractedDocument(pages=pages, source="pdf")


def render_pdf_pages(data: bytes, *, limit: int = MAX_RENDERED_PAGES) -> list[bytes]:
    """Rend les pages en PNG pour le modele vision."""
    images: list[bytes] = []
    with fitz.open(stream=data, filetype="pdf") as document:
        for page in document[:limit]:
            pixmap = page.get_pixmap(dpi=RENDER_DPI)
            images.append(pixmap.tobytes("png"))
    return images


# --- DOCX ------------------------------------------------------------------
def extract_docx(data: bytes) -> ExtractedDocument:
    """Extrait un DOCX. Pas de coordonnees : le format n'a pas de mise en page fixe.

    Les tableaux sont parcourus explicitement — beaucoup de CV y placent les
    competences et les langues, et `document.paragraphs` les ignore.
    """
    try:
        document = DocxDocument(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"DOCX illisible : {exc}") from exc

    lines = [para.text.strip() for para in document.paragraphs if para.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    text = "\n".join(lines)
    return ExtractedDocument(
        pages=[Page(number=1, text=text, words=[])], source="docx"
    )


def extract(data: bytes, filename: str) -> ExtractedDocument:
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if suffix == "pdf":
        return extract_pdf(data)
    if suffix == "docx":
        return extract_docx(data)
    raise ExtractionError(f"Format non pris en charge : .{suffix}")
