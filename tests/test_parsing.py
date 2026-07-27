"""Tests du pipeline d'extraction.

Aucun serveur d'inference n'est requis : l'appel modele est remplace par une
charge utile fixe. Ce qui est teste ici, c'est tout ce qui l'entoure —
extraction, diagnostic de mise en page, ancrage des preuves, persistance.
"""

from __future__ import annotations

import datetime as dt
import functools

import fitz
import pytest
from django.core.exceptions import ValidationError
from django.core.files.uploadedfile import SimpleUploadedFile

from apps.candidates.models import Candidate, CVDocument
from apps.parsing import evidence as evidence_module
from apps.parsing import extractors, pipeline, quality
from apps.parsing.evidence import EvidenceResolver
from apps.parsing.services import ingest, validate_upload

LOREM = (
    "Ingenieur logiciel specialise en developpement backend Python. "
    "Conception d'APIs REST avec Django et Django REST Framework. "
    "Mise en place de pipelines de donnees et integration de modeles de langage. "
    "Encadrement technique d'une equipe de quatre developpeurs sur deux ans. "
)


# --- Fabrication de PDF de test --------------------------------------------
# Mise en cache : PyMuPDF ecrit un identifiant de document different a chaque
# generation. Sans cela, deux appels produiraient deux hashs distincts et le
# test de deduplication par hash de contenu n'aurait aucun sens.
@functools.lru_cache(maxsize=1)
def _single_column_pdf() -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_textbox(fitz.Rect(50, 50, 545, 780), LOREM * 6, fontsize=10)
    data = document.tobytes()
    document.close()
    return data


@functools.lru_cache(maxsize=1)
def _two_column_pdf() -> bytes:
    document = fitz.open()
    page = document.new_page()
    # Colonne gauche 50-260, colonne droite 340-545 : couloir vide de 80 points.
    page.insert_textbox(fitz.Rect(50, 50, 260, 780), LOREM * 4, fontsize=9)
    page.insert_textbox(fitz.Rect(340, 50, 545, 780), LOREM * 4, fontsize=9)
    data = document.tobytes()
    document.close()
    return data


def _empty_pdf() -> bytes:
    document = fitz.open()
    document.new_page()
    data = document.tobytes()
    document.close()
    return data


# --- Extraction ------------------------------------------------------------
def test_pdf_extraction_produces_words_with_offsets():
    extracted = extractors.extract(_single_column_pdf(), "cv.pdf")
    assert extracted.page_count == 1
    page = extracted.pages[0]
    assert page.words

    # L'offset de chaque mot doit pointer exactement sur ce mot dans le texte.
    for word in page.words[:50]:
        assert page.text[word.char_start : word.char_end] == word.text


def test_pdf_extraction_reads_content():
    extracted = extractors.extract(_single_column_pdf(), "cv.pdf")
    assert "Django" in extracted.full_text
    assert extracted.char_count > 500


def test_docx_extraction_reads_paragraphs_and_tables():
    from io import BytesIO

    from docx import Document as DocxDocument

    document = DocxDocument()
    document.add_paragraph("Badr Sahraoui")
    document.add_paragraph("Developpeur backend Python")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Langues"
    table.rows[0].cells[1].text = "Francais, Anglais"
    buffer = BytesIO()
    document.save(buffer)

    extracted = extractors.extract(buffer.getvalue(), "cv.docx")
    assert extracted.source == "docx"
    assert "Badr Sahraoui" in extracted.full_text
    assert "Francais, Anglais" in extracted.full_text  # le tableau est bien lu


def test_unsupported_format_raises():
    with pytest.raises(extractors.ExtractionError):
        extractors.extract(b"x", "cv.txt")


def test_render_pages_returns_png():
    images = extractors.render_pdf_pages(_single_column_pdf())
    assert len(images) == 1
    assert images[0].startswith(b"\x89PNG")


# --- Diagnostic de qualite -------------------------------------------------
def test_single_column_is_not_flagged():
    report = quality.assess(extractors.extract(_single_column_pdf(), "cv.pdf"))
    assert report.has_text_layer
    assert not report.looks_scanned
    assert not report.is_multi_column
    assert not report.needs_vision


def test_two_column_layout_is_detected():
    report = quality.assess(extractors.extract(_two_column_pdf(), "cv.pdf"))
    assert report.multi_column_pages == [1]
    assert report.needs_vision  # bascule vers le modele vision


def test_empty_pdf_looks_scanned():
    report = quality.assess(extractors.extract(_empty_pdf(), "cv.pdf"))
    assert not report.has_text_layer
    assert report.looks_scanned
    assert report.needs_vision


# --- Ancrage des preuves ---------------------------------------------------
@pytest.fixture
def resolver() -> EvidenceResolver:
    return EvidenceResolver(extractors.extract(_single_column_pdf(), "cv.pdf"))


def test_exact_quote_is_anchored(resolver):
    resolved = resolver.resolve("Conception d'APIs REST avec Django")
    assert resolved is not None
    assert resolved.verified
    assert resolved.ratio == 1.0
    assert resolved.page == 1
    assert resolved.bbox is not None and len(resolved.bbox) == 4


def test_anchoring_ignores_case_and_accents(resolver):
    resolved = resolver.resolve("CONCEPTION D'APIS REST AVEC DJANGO")
    assert resolved is not None and resolved.verified


def test_partial_quote_is_anchored_with_lower_ratio(resolver):
    resolved = resolver.resolve("Conception d'APIs REST avec Symfony et Laravel")
    assert resolved is not None
    assert resolved.ratio < 1.0


def test_invented_quote_is_rejected(resolver):
    assert resolver.resolve("Certification AWS Solutions Architect Professional") is None


def test_short_quote_is_accepted_with_reduced_confidence(resolver):
    """Une citation breve n'est pas une hallucination.

    Sur un vrai CV, Qwen a cite « Python », « Django », « Docker » : des
    intitules exacts, presents dans le document, mais peu bavards. Les traiter
    comme des citations inventees faisait remonter trois competences reelles en
    « non etayees ».
    """
    resolved = resolver.resolve("Django")
    assert resolved is not None
    assert resolved.verified
    assert resolved.ratio == evidence_module.SHORT_QUOTE_RATIO
    assert resolved.ratio < 1.0  # moins probante qu'une phrase entiere


def test_short_quote_requires_a_whole_word(resolver):
    """« SQL » ne doit pas se valider sur « PostgreSQL »."""
    assert resolver.resolve("Symfony") is None


def test_quote_below_three_characters_is_rejected(resolver):
    assert resolver.resolve("de") is None
    assert resolver.resolve("") is None


# --- Utilitaires du pipeline -----------------------------------------------
@pytest.mark.parametrize(
    ("value", "expected"),
    [
        ("2023-05", dt.date(2023, 5, 1)),
        ("2023", dt.date(2023, 1, 1)),
        ("2023-13", None),
        ("hier", None),
        ("", None),
        (None, None),
    ],
)
def test_month_parsing(value, expected):
    assert pipeline._parse_month(value) == expected


@pytest.mark.parametrize(
    ("value", "expected"), [(4, 3), (5, 5), (7, 8), (2, 1), ("x", 0), (None, 0)]
)
def test_education_level_is_snapped_to_valid_choice(value, expected):
    assert pipeline._closest_education_level(value) == expected


def test_url_cleaning():
    assert pipeline._clean_url("linkedin.com/in/badr") == "https://linkedin.com/in/badr"
    assert pipeline._clean_url("https://github.com/badr") == "https://github.com/badr"
    assert pipeline._clean_url("") == ""


def test_overlapping_experiences_are_not_counted_twice(db):
    from apps.candidates.models import Experience

    candidate = Candidate.objects.create(full_name="Test")
    # Deux missions menees en parallele sur 2020-2022 : 2 ans, pas 4.
    Experience.objects.create(
        candidate=candidate, title="A",
        start_date=dt.date(2020, 1, 1), end_date=dt.date(2022, 1, 1),
    )
    Experience.objects.create(
        candidate=candidate, title="B",
        start_date=dt.date(2020, 6, 1), end_date=dt.date(2022, 1, 1),
    )
    assert pipeline._total_experience_years(candidate) == pytest.approx(2.0, abs=0.05)


def test_consecutive_experiences_are_summed(db):
    from apps.candidates.models import Experience

    candidate = Candidate.objects.create(full_name="Test")
    Experience.objects.create(
        candidate=candidate, title="A",
        start_date=dt.date(2018, 1, 1), end_date=dt.date(2020, 1, 1),
    )
    Experience.objects.create(
        candidate=candidate, title="B",
        start_date=dt.date(2020, 1, 1), end_date=dt.date(2023, 1, 1),
    )
    assert pipeline._total_experience_years(candidate) == pytest.approx(5.0, abs=0.05)


# --- Depot ------------------------------------------------------------------
def test_upload_rejects_unknown_extension():
    with pytest.raises(ValidationError):
        validate_upload(SimpleUploadedFile("cv.txt", b"x"))


def test_upload_rejects_oversized_file(settings):
    settings.MAX_CV_SIZE_BYTES = 10
    with pytest.raises(ValidationError):
        validate_upload(SimpleUploadedFile("cv.pdf", b"x" * 100))


# --- Pipeline complet -------------------------------------------------------
PAYLOAD = {
    "identity": {
        "full_name": "Badr Sahraoui",
        "email": "Badr.Sahraoui@example.com",
        "phone": "+212 600 000 000",
        "linkedin": "linkedin.com/in/badr",
        "github": "",
        "location": "Casablanca",
        "headline": "Ingenieur logiciel backend",
    },
    "skills": [
        {
            "name": "Django",
            "years": 3,
            "last_used_year": 2026,
            "evidence": "Conception d'APIs REST avec Django",
        },
        {
            "name": "Python",
            "years": 4,
            "last_used_year": 2026,
            "evidence": "developpement backend Python",
        },
        {
            # Citation absente du document : doit finir non etayee.
            "name": "Kubernetes",
            "years": 2,
            "last_used_year": 2025,
            "evidence": "Administration d'un cluster Kubernetes en production",
        },
    ],
    "experiences": [
        {
            "title": "Ingenieur backend",
            "company": "Exemple SARL",
            "location": "Casablanca",
            "start_date": "2022-01",
            "end_date": "2024-01",
            "description": "APIs Django",
            "evidence": "Encadrement technique d'une equipe",
        }
    ],
    "education": [
        {
            "degree": "Master informatique",
            "field_of_study": "Genie logiciel",
            "institution": "Universite",
            "level": 5,
            "graduation_year": 2021,
            "evidence": "Ingenieur logiciel specialise",
        }
    ],
    "languages": [
        {"language": "Francais", "level": "NAT", "evidence": "Conception d'APIs REST"},
        {"language": "Anglais", "level": "XX", "evidence": ""},
    ],
    "certifications": [],
}


@pytest.fixture
def parsed_document(db, settings, tmp_path, monkeypatch):
    settings.MEDIA_ROOT = tmp_path
    monkeypatch.setattr(
        pipeline,
        "_structure",
        lambda document, data, extracted, report: (PAYLOAD, CVDocument.Method.TEXT),
    )
    upload = SimpleUploadedFile("cv.pdf", _single_column_pdf(), content_type="application/pdf")
    document, created = ingest(upload)
    assert created
    document.refresh_from_db()
    return document


def test_pipeline_completes(parsed_document):
    assert parsed_document.status == CVDocument.Status.DONE
    assert parsed_document.method == CVDocument.Method.TEXT
    assert parsed_document.page_count == 1
    assert parsed_document.extraction_seconds is not None
    assert parsed_document.quality["needs_vision"] is False


def test_pipeline_creates_candidate(parsed_document):
    candidate = parsed_document.candidate
    assert candidate.full_name == "Badr Sahraoui"
    assert candidate.email == "badr.sahraoui@example.com"  # normalise en minuscules
    assert candidate.linkedin_url == "https://linkedin.com/in/badr"
    assert candidate.headline == "Ingenieur logiciel backend"
    assert candidate.highest_education == 5
    assert candidate.total_experience_years == pytest.approx(2.0, abs=0.05)


def test_pipeline_anchors_verified_evidence(parsed_document):
    candidate = parsed_document.candidate
    django = candidate.skills.get(normalized_name="django")
    assert django.evidence is not None
    assert django.evidence.verified
    assert django.evidence.bbox is not None
    assert django.confidence == 1.0


def test_invented_evidence_is_flagged_not_dropped(parsed_document):
    """Une citation introuvable ne fait pas disparaitre la donnee : elle la degrade."""
    kubernetes = parsed_document.candidate.skills.get(normalized_name="kubernetes")
    assert kubernetes.evidence is not None
    assert not kubernetes.evidence.verified
    assert kubernetes.confidence < 1.0


def test_invalid_language_level_falls_back(parsed_document):
    anglais = parsed_document.candidate.languages.get(language="Anglais")
    assert anglais.level == "B1"
    francais = parsed_document.candidate.languages.get(language="Francais")
    assert francais.level == "NAT"


def test_same_file_is_not_reprocessed(parsed_document, monkeypatch):
    """Le hash du contenu sert de cache : redeposer le meme CV ne relance rien."""
    calls = []
    monkeypatch.setattr(
        pipeline, "_structure",
        lambda *args: calls.append(1) or (PAYLOAD, CVDocument.Method.TEXT),
    )
    upload = SimpleUploadedFile("copie.pdf", _single_column_pdf(), content_type="application/pdf")
    document, created = ingest(upload)

    assert not created
    assert document.pk == parsed_document.pk
    assert calls == []
    assert Candidate.objects.count() == 1


def test_audit_entry_is_written(parsed_document):
    from apps.core.models import AuditLog

    actions = set(AuditLog.objects.values_list("action", flat=True))
    assert AuditLog.Action.CV_UPLOADED in actions
    assert AuditLog.Action.CV_PARSED in actions

    entry = AuditLog.objects.get(action=AuditLog.Action.CV_PARSED)
    assert entry.metadata["evidence_total"] > 0
    assert entry.metadata["method"] == CVDocument.Method.TEXT


def test_failure_is_recorded_on_document(db, settings, tmp_path, monkeypatch):
    settings.MEDIA_ROOT = tmp_path

    def boom(*args, **kwargs):
        raise pipeline.ParsingError("serveur d'inference injoignable")

    monkeypatch.setattr(pipeline, "_structure", boom)
    upload = SimpleUploadedFile("ko.pdf", _single_column_pdf(), content_type="application/pdf")

    with pytest.raises(pipeline.ParsingError):
        ingest(upload)

    document = CVDocument.objects.get(original_filename="ko.pdf")
    assert document.status == CVDocument.Status.FAILED
    assert "injoignable" in document.error
